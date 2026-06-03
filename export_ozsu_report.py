"""
Xuất báo cáo Word: Phân tích Hệ thống theo khung lý thuyết Özsu & Valduriez
"Principles of Distributed Database Systems"
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup: A4, margin 2.5cm ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(3)

# ── Helper functions ──
def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={0:16, 1:14, 2:13}.get(level, 12), bold=True, color=(0,0,0))
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 0 else WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_para(text, bold=False, italic=False, indent_cm=0, align='justify'):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p

def add_mixed_para(parts, indent_cm=0):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    if indent_level:
        p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
    return p

def add_table_grid(headers, data):
    table = doc.add_table(rows=1+len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
    for r, row_data in enumerate(data, start=1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=11)
    doc.add_paragraph()
    return table

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, italic=True)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

# ════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title.add_run('BÁO CÁO PHÂN TÍCH HỆ THỐNG\nCƠ SỞ DỮ LIỆU PHÂN TÁN')
set_run_font(run_t, size=18, bold=True)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = subtitle.add_run('ONLINE BANKING LEDGER — DRIFT DETECTOR\n& DYNAMIC RE-FRAGMENTATION')
set_run_font(run_s, size=14, bold=True, color=(0, 51, 102))

doc.add_paragraph()
doc.add_paragraph()

framework = doc.add_paragraph()
framework.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = framework.add_run('Khung lý thuyết áp dụng:\nM. Tamer Özsu & Patrick Valduriez\n"Principles of Distributed Database Systems" (4th Edition)')
set_run_font(run_f, size=13, italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════
# MỤC LỤC
# ════════════════════════════════════════════════════
add_heading_styled('MỤC LỤC', level=0)
toc_items = [
    'I. Giới thiệu và Bối cảnh Lý thuyết',
    'II. Thiết kế Phân mảnh (Fragmentation Design)',
    'III. Phân bổ Dữ liệu (Data Allocation)',
    'IV. Mức độ Minh bạch (Transparency Levels)',
    'V. Xử lý Truy vấn Phân tán (Distributed Query Processing)',
    'VI. Quản lý Giao dịch Phân tán (Distributed Transaction Management)',
    'VII. Tái Phân mảnh Động (Dynamic Re-Fragmentation)',
    'VIII. Đánh giá Tính đúng đắn của Phân mảnh',
    'IX. Kết luận'
]
for item in toc_items:
    add_para(item, indent_cm=0.5)

doc.add_page_break()

# ════════════════════════════════════════════════════
# I. GIỚI THIỆU
# ════════════════════════════════════════════════════
add_heading_styled('I. GIỚI THIỆU VÀ BỐI CẢNH LÝ THUYẾT', level=1)

add_para(
    'Theo Özsu và Valduriez (2020), một Cơ sở dữ liệu phân tán (Distributed Database — DDB) '
    'là một tập hợp của nhiều cơ sở dữ liệu có liên quan về mặt logic nhưng được phân bổ trên '
    'một mạng máy tính, và một Hệ quản trị cơ sở dữ liệu phân tán (Distributed DBMS — DDBMS) '
    'là phần mềm quản lý cơ sở dữ liệu phân tán trong khi cung cấp cho người dùng một cơ chế '
    'truy cập minh bạch (transparent access mechanism).'
)

add_para(
    'Hệ thống "Online Banking Ledger — Drift Detector" được xây dựng như một mô phỏng thực tế '
    'của các nguyên lý cốt lõi trong DDBMS theo khung lý thuyết của Özsu và Valduriez. Hệ thống '
    'mô phỏng một ngân hàng trực tuyến có hai chi nhánh địa lý (East Coast và West Coast), mỗi '
    'chi nhánh vận hành một phân vùng cơ sở dữ liệu PostgreSQL độc lập. Một thành phần điều phối '
    'trung tâm (Coordinator) đảm nhận vai trò kết hợp, giám sát và tối ưu hóa dữ liệu phân tán.'
)

add_mixed_para([
    ('Điểm nổi bật của hệ thống là việc mở rộng khung lý thuyết cổ điển của Özsu và Valduriez '
     'bằng cơ chế ', False, False),
    ('Tái phân mảnh động (Dynamic Re-Fragmentation)', True, False),
    (' — một quá trình giám sát liên tục chỉ số ', False, False),
    ('Locality of Reference (LR)', True, True),
    (' và tự động dịch chuyển dữ liệu khi phát hiện sự suy giảm hiệu năng do thay đổi '
     'hành vi người dùng (Fragmentation Drift). Cơ chế này kết hợp các ý tưởng từ '
     'SWORD (giám sát giao dịch dựa trên ngưỡng) và E-Store (lập hồ sơ dữ liệu cấp bản ghi).', False, False)
])

# ════════════════════════════════════════════════════
# II. THIẾT KẾ PHÂN MẢNH
# ════════════════════════════════════════════════════
add_heading_styled('II. THIẾT KẾ PHÂN MẢNH (FRAGMENTATION DESIGN)', level=1)

add_mixed_para([
    ('Theo Chương 3 của Özsu và Valduriez, ', False, False),
    ('phân mảnh (fragmentation)', True, True),
    (' là quá trình chia một quan hệ toàn cục (global relation) thành các đoạn nhỏ hơn '
     '(fragments) để lưu trữ tại các vị trí khác nhau. Có ba phương pháp phân mảnh cơ bản: '
     'phân mảnh ngang (horizontal fragmentation), phân mảnh dọc (vertical fragmentation) '
     'và phân mảnh hỗn hợp (hybrid fragmentation).', False, False)
])

add_heading_styled('2.1. Phân mảnh Ngang Sơ cấp (Primary Horizontal Fragmentation)', level=2)

add_para(
    'Hệ thống áp dụng phương pháp Phân mảnh ngang sơ cấp cho cả hai bảng chính. Theo '
    'Özsu và Valduriez, phân mảnh ngang sơ cấp của một quan hệ R được định nghĩa thông qua '
    'một tập hợp các vị từ chọn lọc (selection predicates) p₁, p₂, ..., pₙ sao cho mỗi '
    'phân mảnh Rᵢ = σ(pᵢ)(R).'
)

add_para('a) Phân mảnh bảng Customers:', bold=True, indent_cm=0.5)
add_para(
    'Vị từ phân mảnh (fragmentation predicate) được xác định dựa trên thuộc tính '
    'homebranchid — chi nhánh đăng ký ban đầu của khách hàng:',
    indent_cm=0.5
)
add_formula('F₁ = σ(homebranchid = \'A\')(customers) → Lưu tại Site A')
add_formula('F₂ = σ(homebranchid = \'B\')(customers) → Lưu tại Site B')

add_para(
    'Trong hệ thống, điều này được hiện thực hóa tại init_db.py khi phân phối 5.000 khách hàng '
    'về hai site: khách hàng được sinh ngẫu nhiên với xác suất 50/50 cho mỗi chi nhánh, và chỉ '
    'được INSERT vào site tương ứng.',
    indent_cm=0.5
)

add_para('b) Phân mảnh bảng Transactions:', bold=True, indent_cm=0.5)
add_para(
    'Bảng giao dịch được phân mảnh ngang dựa trên thuộc tính atm_branchid — vị trí vật lý '
    'của ATM nơi giao dịch phát sinh:',
    indent_cm=0.5
)
add_formula('F₃ = σ(atm_branchid = \'A\')(transactions) → Lưu tại Site A')
add_formula('F₄ = σ(atm_branchid = \'B\')(transactions) → Lưu tại Site B')

add_mixed_para([
    ('Lưu ý: ', True, False),
    ('Đây cũng là phân mảnh ngang sơ cấp (không phải phân mảnh ngang dẫn xuất) vì vị từ '
     'phân mảnh (atm_branchid) là thuộc tính nội tại của chính bảng transactions, không phải '
     'được dẫn xuất từ bảng customers thông qua phép nối.', False, False)
], indent_cm=0.5)

add_heading_styled('2.2. Sao chép Dữ liệu (Replication)', level=2)

add_mixed_para([
    ('Theo Özsu và Valduriez, ', False, False),
    ('sao chép đầy đủ (full replication)', True, True),
    (' là chiến lược lưu trữ bản sao hoàn chỉnh của một quan hệ tại tất cả các site. '
     'Trong hệ thống này, bảng ', False, False),
    ('branches', True, False),
    (' (chứa thông tin chi nhánh ngân hàng) được sao chép đầy đủ trên cả hai Site A và Site B. '
     'Điều này là cần thiết vì bảng branches đóng vai trò bảng tham chiếu cho ràng buộc khóa '
     'ngoại (FOREIGN KEY) của cả bảng customers và transactions tại mỗi site.', False, False)
])

# ════════════════════════════════════════════════════
# III. PHÂN BỔ DỮ LIỆU
# ════════════════════════════════════════════════════
add_heading_styled('III. PHÂN BỔ DỮ LIỆU (DATA ALLOCATION)', level=1)

add_para(
    'Theo Chương 3 của Özsu và Valduriez, bài toán phân bổ dữ liệu (Data Allocation Problem) '
    'là quá trình quyết định mỗi phân mảnh sẽ được lưu trữ tại site nào trong mạng. Bài toán '
    'này nhằm tối thiểu hóa chi phí truyền dữ liệu và tối đa hóa hiệu năng truy vấn cục bộ.'
)

add_para('Hệ thống áp dụng chiến lược phân bổ như sau:', bold=True)

add_table_grid(
    ['Quan hệ', 'Chiến lược Phân bổ', 'Mục tiêu'],
    [
        ['branches', 'Sao chép đầy đủ (Full Replication) — tất cả sites',
         'Đảm bảo toàn vẹn tham chiếu (FK) cục bộ'],
        ['customers (F₁)', 'Phân bổ tại Site A (không sao chép)',
         'Tối đa hóa truy cập cục bộ cho Branch A'],
        ['customers (F₂)', 'Phân bổ tại Site B (không sao chép)',
         'Tối đa hóa truy cập cục bộ cho Branch B'],
        ['transactions (F₃)', 'Phân bổ tại Site A (không sao chép)',
         'Ghi trực tiếp tại nơi ATM phát sinh'],
        ['transactions (F₄)', 'Phân bổ tại Site B (không sao chép)',
         'Ghi trực tiếp tại nơi ATM phát sinh'],
    ]
)

add_para(
    'Chiến lược phân bổ không sao chép (non-replicated allocation) cho customers và transactions '
    'tuân theo nguyên tắc "dữ liệu nằm gần nơi sử dụng nhiều nhất" (locality principle) của '
    'Özsu và Valduriez. Tuy nhiên, chiến lược này dẫn đến vấn đề hiệu năng khi hành vi người '
    'dùng thay đổi — đây chính là động lực cho cơ chế tái phân mảnh động được trình bày ở Mục VII.'
)

# ════════════════════════════════════════════════════
# IV. TRANSPARENCY
# ════════════════════════════════════════════════════
add_heading_styled('IV. MỨC ĐỘ MINH BẠCH (TRANSPARENCY LEVELS)', level=1)

add_para(
    'Özsu và Valduriez (Chương 1) xác định ba mức độ minh bạch trong hệ thống DDBMS: '
    'minh bạch phân mảnh (fragmentation transparency), minh bạch vị trí (location transparency) '
    'và minh bạch sao chép (replication transparency). Mỗi mức độ xác định mức ẩn dấu '
    'thông tin phân tán đối với người dùng cuối.'
)

add_table_grid(
    ['Loại Minh bạch', 'Định nghĩa (Özsu & Valduriez)', 'Mức độ đạt được'],
    [
        ['Minh bạch phân mảnh\n(Fragmentation Transparency)',
         'Người dùng không cần biết dữ liệu bị phân mảnh hay không.',
         'Đạt được một phần: Giao diện Dashboard ẩn việc dữ liệu nằm ở site nào '
         'khi xem chi tiết KH. Tuy nhiên, có 2 tab riêng biệt cho Site A và Site B.'],
        ['Minh bạch vị trí\n(Location Transparency)',
         'Người dùng không cần biết dữ liệu vật lý nằm ở đâu.',
         'Đạt được: API endpoints (VD: /api/customer/<id>/transactions) tự động '
         'truy vấn cả 2 sites và trả về kết quả hợp nhất. Người dùng không cần chỉ định site.'],
        ['Minh bạch sao chép\n(Replication Transparency)',
         'Người dùng không cần biết dữ liệu có bản sao hay không.',
         'Đạt được: Bảng branches được sao chép tự động khi init_db.py chạy. '
         'Coordinator xử lý mọi thao tác ghi/đọc minh bạch.'],
    ]
)

add_para(
    'Đặc biệt, chức năng xem lịch sử giao dịch của một khách hàng (API /api/customer/<id>/transactions) '
    'thể hiện rõ nhất tính minh bạch phân mảnh: Coordinator tự động gửi truy vấn đến cả hai site, '
    'hợp nhất kết quả và trả về danh sách hoàn chỉnh cho người dùng mà không yêu cầu họ biết '
    'giao dịch nào nằm ở site nào.'
)

# ════════════════════════════════════════════════════
# V. XỬ LÝ TRUY VẤN PHÂN TÁN
# ════════════════════════════════════════════════════
add_heading_styled('V. XỬ LÝ TRUY VẤN PHÂN TÁN (DISTRIBUTED QUERY PROCESSING)', level=1)

add_para(
    'Theo Chương 7–8 của Özsu và Valduriez, xử lý truy vấn phân tán gồm bốn giai đoạn: '
    '(1) Phân rã truy vấn (Query Decomposition), (2) Cục bộ hóa dữ liệu (Data Localization), '
    '(3) Tối ưu hóa truy vấn toàn cục (Global Query Optimization), và '
    '(4) Tối ưu hóa truy vấn cục bộ (Local Query Optimization).'
)

add_heading_styled('5.1. Phân rã và Cục bộ hóa Truy vấn', level=2)

add_para(
    'Trong hệ thống này, vai trò phân rã truy vấn và cục bộ hóa dữ liệu được thực hiện '
    'bởi Flask Coordinator tại tầng ứng dụng (application-level query decomposition). '
    'Coordinator nhận yêu cầu từ người dùng, xác định các site cần truy vấn, gửi các '
    'truy vấn con (sub-queries) đến từng site qua kết nối psycopg2, và hợp nhất kết quả.'
)

add_para('Ví dụ minh họa với API tính Locality of Reference:', bold=True)

add_code_block(
    '# Bước 1: Thu thập customers từ cả 2 site (Cục bộ hóa)\n'
    'cust_branch = {}\n'
    'for s in [\'a\', \'b\']:\n'
    '    for r in query(s, "SELECT customerid, homebranchid FROM customers"):\n'
    '        cust_branch[r[\'customerid\']] = r[\'homebranchid\']\n\n'
    '# Bước 2: Thu thập transactions từ cả 2 site\n'
    'all_tx = []\n'
    'for s in [\'a\', \'b\']:\n'
    '    rows = query(s, "SELECT customerid, atm_branchid FROM transactions WHERE workload_day=%s", (day,))\n'
    '    all_tx.extend(rows)\n\n'
    '# Bước 3: Thực hiện JOIN tại Application Layer (In-Application Distributed Join)\n'
    'for tx in all_tx:\n'
    '    home = cust_branch.get(tx[\'customerid\'])\n'
    '    key = home + \'_\' + tx[\'atm_branchid\']\n'
    '    stats[key] = stats.get(key, 0) + 1'
)

add_heading_styled('5.2. Kết nối Phân tán (Distributed Join)', level=2)

add_mixed_para([
    ('Özsu và Valduriez phân loại các phương pháp thực hiện phép kết nối phân tán gồm: '
     'Ship-whole, Fetch-as-needed và Semi-join. Hệ thống này sử dụng chiến lược ', False, False),
    ('Ship-whole (Gửi toàn bộ)', True, True),
    (': tất cả dữ liệu liên quan được gửi về Coordinator, phép kết nối được thực hiện '
     'tại tầng ứng dụng (Python dictionary lookup) thay vì tại tầng cơ sở dữ liệu. '
     'Đây là phương pháp đơn giản nhưng hiệu quả cho quy mô dữ liệu của hệ thống '
     '(~5.000 khách hàng, ~36.000 giao dịch).', False, False)
])

add_para(
    'Ưu điểm của phương pháp này là không cần thiết lập cơ chế liên kết cơ sở dữ liệu phức tạp '
    '(như PostgreSQL Foreign Data Wrapper hoặc dblink), giảm thiểu phụ thuộc hạ tầng. '
    'Nhược điểm là toàn bộ dữ liệu phải được truyền về Coordinator, không phù hợp cho '
    'hệ thống có hàng triệu bản ghi.'
)

# ════════════════════════════════════════════════════
# VI. QUẢN LÝ GIAO DỊCH PHÂN TÁN
# ════════════════════════════════════════════════════
add_heading_styled('VI. QUẢN LÝ GIAO DỊCH PHÂN TÁN (DISTRIBUTED TRANSACTION MANAGEMENT)', level=1)

add_para(
    'Theo Chương 10–11 của Özsu và Valduriez, quản lý giao dịch trong môi trường phân tán '
    'phải đảm bảo các tính chất ACID (Atomicity, Consistency, Isolation, Durability) trên '
    'nhiều site. Hai vấn đề cốt lõi là: (1) Kiểm soát đồng thời phân tán '
    '(Distributed Concurrency Control), và (2) Phục hồi phân tán (Distributed Recovery).'
)

add_heading_styled('6.1. Tính Nguyên tử (Atomicity) trong Migration', level=2)

add_para(
    'Thao tác tái phân mảnh (Migration) là một giao dịch phân tán phức tạp liên quan đến '
    'cả hai site. Hệ thống đảm bảo tính nguyên tử thông qua cơ chế Backup-Restore:'
)

add_bullet('Trước khi thực hiện migration, tạo bảng sao lưu (customers_backup, '
           'transactions_backup) trên cả hai site bằng CREATE TABLE ... AS SELECT *.')
add_bullet('Nếu xảy ra lỗi trong quá trình migration, người dùng có thể kích hoạt '
           'chức năng Undo để khôi phục toàn bộ dữ liệu từ bảng backup.')
add_bullet('Cơ chế ON CONFLICT (txid) DO NOTHING được sử dụng để đảm bảo tính '
           'bình đẳng (idempotent) — chạy lại thao tác không gây lỗi trùng khóa.')

add_heading_styled('6.2. Tính Nhất quán (Consistency) — Đồng bộ Sequence', level=2)

add_para(
    'Một thách thức đặc thù trong hệ thống phân tán là đảm bảo tính nhất quán của các '
    'bộ đếm tự tăng (auto-increment sequences). Khi giao dịch được di chuyển giữa các '
    'site với giá trị txid gốc, chuỗi transactions_txid_seq tại site đích có thể bị '
    'lệch so với giá trị txid lớn nhất thực tế.'
)

add_para(
    'Hệ thống giải quyết vấn đề này bằng lệnh đồng bộ sequence sau mỗi thao tác migration:',
)
add_code_block(
    "SELECT setval('transactions_txid_seq', (SELECT MAX(txid) FROM transactions))"
)
add_para(
    'Điều này đảm bảo rằng bất kỳ giao dịch mới nào được tạo sau migration sẽ nhận giá '
    'trị txid tiếp theo hợp lệ, tránh xung đột khóa chính.'
)

add_heading_styled('6.3. Phục hồi Phân tán (Distributed Recovery)', level=2)

add_para(
    'Chức năng Undo Migration thể hiện cơ chế phục hồi phân tán theo mô hình '
    'TRUNCATE — RESTORE:'
)
add_bullet('Bước 1: Kiểm tra sự tồn tại của bảng backup (information_schema.tables).')
add_bullet('Bước 2: TRUNCATE toàn bộ bảng transactions và customers trên từng site (CASCADE).')
add_bullet('Bước 3: INSERT toàn bộ dữ liệu gốc từ bảng backup trở lại bảng chính.')
add_bullet('Bước 4: Thực hiện tuần tự trên cả hai site (Site A trước, Site B sau).')

add_para(
    'Lưu ý rằng hệ thống hiện chưa triển khai giao thức Two-Phase Commit (2PC) theo mô '
    'hình chuẩn của Özsu và Valduriez. Thay vào đó, Coordinator thực hiện các thao tác '
    'tuần tự trên từng site. Đây là sự đánh đổi phù hợp cho hệ thống mô phỏng học thuật '
    'với hai site và không có yêu cầu đồng thời cao.'
)

# ════════════════════════════════════════════════════
# VII. TÁI PHÂN MẢNH ĐỘNG
# ════════════════════════════════════════════════════
add_heading_styled('VII. TÁI PHÂN MẢNH ĐỘNG (DYNAMIC RE-FRAGMENTATION)', level=1)

add_mixed_para([
    ('Özsu và Valduriez (Chương 3) thảo luận về thiết kế phân mảnh như một quyết định thiết '
     'kế thời điểm (design-time decision). Tuy nhiên, trong thực tế, hành vi người dùng '
     'thay đổi theo thời gian dẫn đến hiện tượng ', False, False),
    ('Fragmentation Drift', True, True),
    (' — khi chiến lược phân mảnh ban đầu không còn tối ưu. '
     'Hệ thống mở rộng khung lý thuyết của Özsu và Valduriez bằng cơ chế giám sát '
     'và tái phân mảnh động.', False, False)
])

add_heading_styled('7.1. Chỉ số Locality of Reference (LR)', level=2)

add_para(
    'Chỉ số LR đo lường hiệu quả của phân mảnh hiện tại bằng cách tính tỷ lệ '
    'giao dịch nội vùng (local transactions) trên tổng số giao dịch:'
)
add_formula('LR(i) = |{tx ∈ Tᵢ : tx.atm_branchid = cᵢ.homebranchid}| / |Tᵢ|')

add_para(
    'Trong đó Tᵢ là tập toàn bộ giao dịch của khách hàng i trên tất cả các site (thu '
    'thập bằng Distributed Query), và cᵢ.homebranchid là chi nhánh hiện tại của khách '
    'hàng. Chỉ số LR dao động từ 0 (toàn bộ giao dịch là liên vùng) đến 1 (toàn bộ '
    'giao dịch là nội vùng).'
)

add_heading_styled('7.2. Phát hiện Drift (Drift Detection)', level=2)

add_para(
    'Hệ thống sử dụng ngưỡng LR có thể cấu hình (mặc định LR_THRESHOLD = 0.70) để '
    'xác định trạng thái Drift. Khi LR của một khách hàng giảm xuống dưới ngưỡng, '
    'khách hàng đó được đánh dấu là ứng viên di trú (migration candidate).'
)

add_para('Kịch bản mô phỏng trong hệ thống:', bold=True)

add_table_grid(
    ['Kịch bản', 'Tỷ lệ TX nội vùng', 'LR Trung bình', 'Trạng thái'],
    [
        ['Day 1 (Ban đầu)', '~90%', '≈ 0.90', 'Ổn định (LR > 0.70)'],
        ['Day 30 (Sau drift)', '~35–45%', '≈ 0.35–0.45', 'Drift Detected (LR < 0.70)'],
        ['Sau Re-Fragmentation', '~85–95%', '> 0.85', 'Phục hồi (LR > 0.70)'],
    ]
)

add_heading_styled('7.3. Quy trình Tái phân mảnh', level=2)

add_para(
    'Quy trình tái phân mảnh gồm bốn bước thực hiện tuần tự:'
)

add_para('Bước 1 — Tạo bản sao lưu phòng ngừa (Backup Phase):', bold=True, indent_cm=0.5)
add_para(
    'Tạo bảng customers_backup và transactions_backup trên cả hai site bằng '
    'CREATE TABLE ... AS SELECT *. Đây là cơ chế bảo vệ dữ liệu để hỗ trợ rollback.',
    indent_cm=1
)

add_para('Bước 2 — Di chuyển bản ghi khách hàng (Customer Migration Phase):', bold=True, indent_cm=0.5)
add_para(
    'Chèn bản ghi khách hàng vào site đích với thuộc tính homebranchid được cập nhật '
    'sang chi nhánh mới. Bản ghi gốc tại site nguồn được xóa sau khi chèn thành công.',
    indent_cm=1
)

add_para('Bước 3 — Di chuyển lịch sử giao dịch (Transaction Migration Phase):', bold=True, indent_cm=0.5)
add_para(
    'Toàn bộ giao dịch của khách hàng được di trú sẽ được sao chép sang site đích '
    'bằng INSERT ... ON CONFLICT (txid) DO NOTHING, giữ nguyên giá trị txid gốc. '
    'Giao dịch gốc tại site nguồn được xóa.',
    indent_cm=1
)

add_para('Bước 4 — Đồng bộ Sequence (Sequence Synchronization Phase):', bold=True, indent_cm=0.5)
add_para(
    'Cập nhật bộ đếm tự tăng transactions_txid_seq trên cả hai site bằng '
    'setval() để tránh xung đột khóa chính ở các giao dịch được tạo sau này.',
    indent_cm=1
)

# ════════════════════════════════════════════════════
# VIII. ĐÁNH GIÁ TÍNH ĐÚNG ĐẮN
# ════════════════════════════════════════════════════
add_heading_styled('VIII. ĐÁNH GIÁ TÍNH ĐÚNG ĐẮN CỦA PHÂN MẢNH', level=1)

add_para(
    'Theo Özsu và Valduriez (Chương 3), một phân mảnh đúng đắn phải thỏa mãn ba quy tắc: '
    'Tính đầy đủ (Completeness), Tính tái tạo (Reconstruction) và Tính rời rạc (Disjointness).'
)

add_heading_styled('8.1. Tính Đầy đủ (Completeness)', level=2)

add_mixed_para([
    ('Định nghĩa: ', True, False),
    ('Mỗi bản ghi trong quan hệ toàn cục R phải thuộc ít nhất một phân mảnh Rᵢ.', False, True)
])

add_para(
    'Đối với bảng customers: Mọi khách hàng có homebranchid ∈ {\'A\', \'B\'} (ràng buộc '
    'bởi FOREIGN KEY tham chiếu branches.branchid), do đó mỗi khách hàng chắc chắn thuộc '
    'phân mảnh F₁ hoặc F₂. Tính đầy đủ được đảm bảo.',
    indent_cm=0.5
)

add_para(
    'Đối với bảng transactions: Mọi giao dịch có atm_branchid ∈ {\'A\', \'B\'} (FOREIGN KEY '
    'tham chiếu branches.branchid), do đó mỗi giao dịch chắc chắn thuộc F₃ hoặc F₄. '
    'Tính đầy đủ được đảm bảo.',
    indent_cm=0.5
)

add_heading_styled('8.2. Tính Tái tạo (Reconstruction)', level=2)

add_mixed_para([
    ('Định nghĩa: ', True, False),
    ('Quan hệ toàn cục R có thể được tái tạo từ các phân mảnh: R = R₁ ∪ R₂ ∪ ... ∪ Rₙ.', False, True)
])

add_para(
    'Với phân mảnh ngang, phép tái tạo là phép hợp (UNION). Bảng customers toàn cục có thể '
    'được tái tạo bằng: customers = F₁ ∪ F₂ (tương đương SQL: SELECT * FROM customers@site_a '
    'UNION ALL SELECT * FROM customers@site_b). Tương tự, bảng transactions toàn cục: '
    'transactions = F₃ ∪ F₄. Coordinator thực hiện phép tái tạo này trong Python khi cần '
    'dữ liệu tổng hợp (ví dụ: hàm migration_candidates() thu thập all_custs từ cả hai site).',
    indent_cm=0.5
)

add_heading_styled('8.3. Tính Rời rạc (Disjointness)', level=2)

add_mixed_para([
    ('Định nghĩa: ', True, False),
    ('Các phân mảnh không có phần tử chung: Rᵢ ∩ Rⱼ = ∅ với mọi i ≠ j.', False, True)
])

add_para(
    'Tính rời rạc được đảm bảo bởi cơ chế lưu trữ vật lý: mỗi bản ghi khách hàng chỉ '
    'tồn tại tại đúng một site (xác định bởi homebranchid). Tương tự, mỗi giao dịch chỉ '
    'nằm tại đúng một site (xác định bởi atm_branchid). Khi thực hiện migration, hệ thống '
    'luôn INSERT vào site đích trước, sau đó DELETE khỏi site nguồn, đảm bảo không có bản '
    'ghi trùng lặp giữa hai site tại bất kỳ thời điểm nào.',
    indent_cm=0.5
)

add_para(
    'Cơ chế ON CONFLICT (txid) DO NOTHING là biện pháp phòng vệ bổ sung: nếu vì bất kỳ '
    'lý do nào mà bản ghi đã tồn tại tại site đích (ví dụ do chạy migration hai lần), '
    'hệ thống sẽ bỏ qua thay vì gây lỗi, duy trì tính rời rạc.',
    indent_cm=0.5
)

# ════════════════════════════════════════════════════
# IX. KẾT LUẬN
# ════════════════════════════════════════════════════
add_heading_styled('IX. KẾT LUẬN', level=1)

add_para(
    'Hệ thống "Online Banking Ledger — Drift Detector" đã triển khai thành công các '
    'nguyên lý cốt lõi của cơ sở dữ liệu phân tán theo khung lý thuyết của Özsu và '
    'Valduriez, bao gồm:'
)

add_bullet('Phân mảnh ngang sơ cấp (Primary Horizontal Fragmentation) cho bảng customers '
           'và transactions dựa trên các vị từ phân mảnh rõ ràng.')
add_bullet('Sao chép đầy đủ (Full Replication) cho bảng tham chiếu branches nhằm đảm bảo '
           'toàn vẹn khóa ngoại cục bộ.')
add_bullet('Phân bổ dữ liệu (Data Allocation) theo nguyên tắc tối đa hóa truy cập cục bộ.')
add_bullet('Xử lý truy vấn phân tán (Distributed Query Processing) tại tầng ứng dụng với '
           'chiến lược Ship-whole cho phép kết nối phân tán.')
add_bullet('Quản lý giao dịch phân tán với cơ chế Backup-Restore đảm bảo khả năng phục hồi.')
add_bullet('Đảm bảo ba quy tắc đúng đắn: Completeness, Reconstruction và Disjointness.')

add_para(
    'Đồng thời, hệ thống đã mở rộng khung lý thuyết cổ điển bằng cơ chế Tái phân mảnh '
    'động (Dynamic Re-Fragmentation) dựa trên chỉ số Locality of Reference — một đóng góp '
    'thực tiễn quan trọng giải quyết bài toán mà Özsu và Valduriez đặt ra: "Điều gì xảy '
    'ra khi giả định thiết kế ban đầu không còn đúng?"'
)

doc.add_paragraph()

# Tài liệu tham khảo
add_heading_styled('TÀI LIỆU THAM KHẢO', level=1)
refs = [
    'Özsu, M. T., & Valduriez, P. (2020). Principles of Distributed Database Systems (4th ed.). Springer.',
    'Tatarowicz, A., Curino, C., Jones, E. P. C., & Madden, S. (2012). Lookup Tables: Fine-Grained Partitioning for Distributed Databases. ICDE.',
    'Quamar, A., Kumar, K. A., & Deshpande, A. (2013). SWORD: Scalable Workload-Aware Data Placement for Transactional Workloads. EDBT.',
    'Taft, R., et al. (2014). E-Store: Fine-Grained Elastic Partitioning for Distributed Transaction Processing Systems. VLDB.',
]
for i, ref in enumerate(refs, 1):
    add_para(f'[{i}] {ref}', indent_cm=0.5)

# ── Save ──
output_path = r'c:\Users\A.Long\OneDrive\Desktop\DDB_Project\Bao_Cao_Ozsu_Valduriez.docx'
doc.save(output_path)
print(f'✅ Đã xuất file thành công: {output_path}')
